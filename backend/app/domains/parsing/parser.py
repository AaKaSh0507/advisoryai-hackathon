import io
import re
import time
from uuid import UUID

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from backend.app.domains.parsing.schemas import (
    BlockType,
    DocumentBlock,
    DocumentMetadata,
    HeaderFooterBlock,
    HeadingBlock,
    ListBlock,
    ListItem,
    PageBreakBlock,
    ParagraphBlock,
    ParsedDocument,
    ParsingError,
    ParsingResult,
    TableBlock,
    TableCell,
    TableRow,
    TextRun,
    generate_block_id,
    generate_content_hash,
)
from backend.app.infrastructure.datetime_utils import utc_now
from backend.app.logging_config import get_logger

logger = get_logger("app.domains.parsing.parser")


class WordDocumentParser:
    HEADING_PATTERNS = [
        re.compile(r"^Heading\s*(\d)$", re.IGNORECASE),
        re.compile(r"^Title$", re.IGNORECASE),
        re.compile(r"^Subtitle$", re.IGNORECASE),
    ]

    LIST_STYLE_PATTERNS = [
        re.compile(r"List", re.IGNORECASE),
        re.compile(r"Bullet", re.IGNORECASE),
    ]

    def __init__(self):
        self._sequence = 0
        self._list_buffer: list[ListItem] = []
        self._current_list_style: str | None = None

    def parse(
        self,
        content: bytes,
        template_id: UUID,
        template_version_id: UUID,
        version_number: int,
    ) -> ParsingResult:
        start_time = time.time()
        errors: list[ParsingError] = []
        warnings: list[str] = []
        self._sequence = 0
        self._list_buffer = []
        self._current_list_style = None

        try:
            content_hash = generate_content_hash(content)
            doc = Document(io.BytesIO(content))
            metadata = self._extract_metadata(doc)
            blocks = self._parse_body(doc, errors, warnings)
            headers, footers = self._parse_headers_footers(doc, errors, warnings)
            parsed_doc = ParsedDocument(
                template_version_id=template_version_id,
                template_id=template_id,
                version_number=version_number,
                content_hash=content_hash,
                parsed_at=utc_now(),
                metadata=metadata,
                blocks=blocks,
                headers=headers,
                footers=footers,
            )
            parsed_doc.compute_statistics()
            end_time = time.time()
            duration_ms = (end_time - start_time) * 1000
            logger.info(
                f"Parsed document successfully: {parsed_doc.total_blocks} blocks "
                f"({parsed_doc.heading_count} headings, {parsed_doc.paragraph_count} paragraphs, "
                f"{parsed_doc.table_count} tables, {parsed_doc.list_count} lists) in {duration_ms:.2f}ms"
            )

            return ParsingResult(
                success=True,
                document=parsed_doc,
                errors=errors,
                warnings=warnings,
                parse_duration_ms=duration_ms,
                total_duration_ms=duration_ms,
            )

        except Exception as e:
            logger.error(f"Failed to parse document: {e}", exc_info=True)
            errors.append(
                ParsingError(
                    error_type="parse_failure",
                    message=str(e),
                    recoverable=False,
                )
            )

            end_time = time.time()
            duration_ms = (end_time - start_time) * 1000

            return ParsingResult(
                success=False,
                errors=errors,
                warnings=warnings,
                total_duration_ms=duration_ms,
            )

    def _extract_metadata(self, doc: DocxDocument) -> DocumentMetadata:
        props = doc.core_properties
        return DocumentMetadata(
            title=props.title,
            author=props.author,
            subject=props.subject,
            keywords=props.keywords,
            created=props.created,
            modified=props.modified,
            revision=props.revision,
        )

    def _parse_body(
        self,
        doc: DocxDocument,
        errors: list[ParsingError],
        warnings: list[str],
    ) -> list[DocumentBlock]:
        blocks: list[DocumentBlock] = []
        body = doc.element.body
        for element in body.iterchildren():
            try:
                if isinstance(element, CT_P):
                    paragraph = Paragraph(element, doc)
                    block = self._parse_paragraph(paragraph, errors, warnings)
                    if block:
                        if self._list_buffer and not self._is_list_item(paragraph):
                            list_block = self._flush_list_buffer()
                            if list_block:
                                blocks.append(list_block)
                        if self._is_list_item(paragraph):
                            self._add_to_list_buffer(paragraph)
                        else:
                            blocks.append(block)

                elif isinstance(element, CT_Tbl):
                    if self._list_buffer:
                        list_block = self._flush_list_buffer()
                        if list_block:
                            blocks.append(list_block)
                    table = Table(element, doc)
                    block = self._parse_table(table, errors, warnings)
                    if block:
                        blocks.append(block)

            except Exception as e:
                logger.warning(f"Error parsing element at sequence {self._sequence}: {e}")
                errors.append(
                    ParsingError(
                        error_type="element_parse_error",
                        message=str(e),
                        block_index=self._sequence,
                        recoverable=True,
                    )
                )
        if self._list_buffer:
            list_block = self._flush_list_buffer()
            if list_block:
                blocks.append(list_block)

        return blocks

    def _parse_paragraph(
        self,
        para: Paragraph,
        _errors: list[ParsingError],
        _warnings: list[str],
    ) -> DocumentBlock | None:
        text = para.text.strip()
        if not text and not para.runs:
            return None
        if self._has_page_break(para):
            self._sequence += 1
            return PageBreakBlock(
                block_id=generate_block_id("page_break", self._sequence),
                sequence=self._sequence,
            )
        runs = self._extract_runs(para)
        heading_level = self._get_heading_level(para)
        self._sequence += 1
        if heading_level:
            return HeadingBlock(
                block_id=generate_block_id("heading", self._sequence, text[:50]),
                sequence=self._sequence,
                level=heading_level,
                runs=runs,
                alignment=self._get_alignment(para),
                style_name=para.style.name if para.style else None,
            )
        else:
            return ParagraphBlock(
                block_id=generate_block_id("paragraph", self._sequence, text[:50]),
                sequence=self._sequence,
                runs=runs,
                alignment=self._get_alignment(para),
                indent_left=self._get_indent(para, "left"),
                indent_right=self._get_indent(para, "right"),
                indent_first_line=self._get_indent(para, "first_line"),
                spacing_before=self._get_spacing(para, "before"),
                spacing_after=self._get_spacing(para, "after"),
                style_name=para.style.name if para.style else None,
            )

    def _extract_runs(self, para: Paragraph) -> list[TextRun]:
        runs: list[TextRun] = []
        for run in para.runs:
            if not run.text:
                continue
            text_run = TextRun(
                text=run.text,
                bold=run.bold or False,
                italic=run.italic or False,
                underline=bool(run.underline),
                strike=run.font.strike or False if run.font else False,
                font_name=run.font.name if run.font else None,
                font_size=run.font.size.pt if run.font and run.font.size else None,
            )
            if run.font and run.font.color and run.font.color.rgb:
                text_run.color = str(run.font.color.rgb)
            runs.append(text_run)

        return runs

    def _get_heading_level(self, para: Paragraph) -> int | None:
        if not para.style:
            return None
        style_name = para.style.name
        for pattern in self.HEADING_PATTERNS:
            match = pattern.match(style_name)
            if match:
                if "Title" in style_name:
                    return 1
                if "Subtitle" in style_name:
                    return 2
                try:
                    return int(match.group(1))
                except (IndexError, ValueError):
                    return 1
        if hasattr(para.style, "paragraph_format") and para.style.paragraph_format:
            outline_level = para.style.paragraph_format.outline_level
            if outline_level is not None and int(outline_level) < 9:
                return int(outline_level) + 1

        return None

    def _is_list_item(self, para: Paragraph) -> bool:
        pPr = para._p.pPr
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                return True
        if para.style:
            for pattern in self.LIST_STYLE_PATTERNS:
                if pattern.search(para.style.name):
                    return True
        return False

    def _add_to_list_buffer(self, para: Paragraph) -> None:
        runs = self._extract_runs(para)
        text = para.text.strip()
        level = 0
        pPr = para._p.pPr
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                ilvl = numPr.find(qn("w:ilvl"))
                if ilvl is not None:
                    try:
                        level = int(ilvl.get(qn("w:val")))
                    except (ValueError, TypeError):
                        level = 0
        style_name = para.style.name if para.style else ""
        if not self._current_list_style:
            if "bullet" in style_name.lower() or "•" in text:
                self._current_list_style = "bullet"
            else:
                self._current_list_style = "numbered"
        item = ListItem(
            item_id=generate_block_id("list_item", len(self._list_buffer), text[:30]),
            level=level,
            runs=runs,
        )
        self._list_buffer.append(item)

    def _flush_list_buffer(self) -> ListBlock | None:
        if not self._list_buffer:
            return None
        self._sequence += 1
        list_block = ListBlock(
            block_id=generate_block_id("list", self._sequence),
            sequence=self._sequence,
            list_type=self._current_list_style or "bullet",
            items=self._list_buffer.copy(),
        )
        self._list_buffer = []
        self._current_list_style = None
        return list_block

    def _parse_table(
        self,
        table: Table,
        errors: list[ParsingError],
        warnings: list[str],
    ) -> TableBlock | None:
        self._sequence += 1
        rows: list[TableRow] = []
        max_cols = 0
        for row_idx, row in enumerate(table.rows):
            cells: list[TableCell] = []
            for col_idx, cell in enumerate(row.cells):
                cell_content = self._parse_cell_content(cell, errors, warnings)
                table_cell = TableCell(
                    cell_id=generate_block_id("cell", row_idx * 100 + col_idx),
                    row_index=row_idx,
                    col_index=col_idx,
                    content=cell_content,
                )
                cells.append(table_cell)
            max_cols = max(max_cols, len(cells))
            table_row = TableRow(
                row_id=generate_block_id("row", row_idx),
                row_index=row_idx,
                cells=cells,
                is_header=row_idx == 0,
            )
            rows.append(table_row)

        return TableBlock(
            block_id=generate_block_id("table", self._sequence),
            sequence=self._sequence,
            rows=rows,
            column_count=max_cols,
        )

    def _parse_cell_content(
        self,
        cell: _Cell,
        _errors: list[ParsingError],
        _warnings: list[str],
    ) -> list[DocumentBlock]:
        blocks: list[DocumentBlock] = []
        for para in cell.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            runs = self._extract_runs(para)
            block = ParagraphBlock(
                block_id=generate_block_id("cell_para", len(blocks), text[:30]),
                sequence=len(blocks),
                runs=runs,
                alignment=self._get_alignment(para),
                style_name=para.style.name if para.style else None,
            )
            blocks.append(block)
        return blocks

    def _parse_headers_footers(
        self,
        doc: DocxDocument,
        _errors: list[ParsingError],
        _warnings: list[str],
    ) -> tuple[list[HeaderFooterBlock], list[HeaderFooterBlock]]:
        headers: list[HeaderFooterBlock] = []
        footers: list[HeaderFooterBlock] = []
        for section in doc.sections:
            for header_type, header in [
                ("default", section.header),
                ("first", section.first_page_header),
                ("even", section.even_page_header),
            ]:
                if header and header.paragraphs:
                    content: list[ParagraphBlock] = []
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            runs = self._extract_runs(para)
                            content.append(
                                ParagraphBlock(
                                    block_id=generate_block_id(
                                        "header_para", len(content), text[:30]
                                    ),
                                    sequence=len(content),
                                    runs=runs,
                                )
                            )

                    if content:
                        self._sequence += 1
                        headers.append(
                            HeaderFooterBlock(
                                block_type=BlockType.HEADER,
                                block_id=generate_block_id("header", self._sequence),
                                sequence=self._sequence,
                                header_footer_type=header_type,
                                content=content,
                            )
                        )

            for footer_type, footer in [
                ("default", section.footer),
                ("first", section.first_page_footer),
                ("even", section.even_page_footer),
            ]:
                if footer and footer.paragraphs:
                    content = []
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            runs = self._extract_runs(para)
                            content.append(
                                ParagraphBlock(
                                    block_id=generate_block_id(
                                        "footer_para", len(content), text[:30]
                                    ),
                                    sequence=len(content),
                                    runs=runs,
                                )
                            )

                    if content:
                        self._sequence += 1
                        footers.append(
                            HeaderFooterBlock(
                                block_type=BlockType.FOOTER,
                                block_id=generate_block_id("footer", self._sequence),
                                sequence=self._sequence,
                                header_footer_type=footer_type,
                                content=content,
                            )
                        )

        return headers, footers

    def _has_page_break(self, para: Paragraph) -> bool:
        for run in para.runs:
            if run._r.xml and "w:br" in run._r.xml and 'w:type="page"' in run._r.xml:
                return True
        return False

    def _get_alignment(self, para: Paragraph) -> str | None:
        if para.alignment is None:
            return None
        return str(para.alignment).replace("WD_PARAGRAPH_ALIGNMENT.", "").lower()

    def _get_indent(self, para: Paragraph, indent_type: str) -> float | None:
        pf = para.paragraph_format
        if not pf:
            return None

        if indent_type == "left" and pf.left_indent:
            return float(pf.left_indent.pt)
        elif indent_type == "right" and pf.right_indent:
            return float(pf.right_indent.pt)
        elif indent_type == "first_line" and pf.first_line_indent:
            return float(pf.first_line_indent.pt)

        return None

    def _get_spacing(self, para: Paragraph, spacing_type: str) -> float | None:
        pf = para.paragraph_format
        if not pf:
            return None

        if spacing_type == "before" and pf.space_before:
            return float(pf.space_before.pt)
        elif spacing_type == "after" and pf.space_after:
            return float(pf.space_after.pt)

        return None
