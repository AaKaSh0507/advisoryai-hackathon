import io
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Twips

from backend.app.domains.rendering.schemas import RenderErrorCode, RenderingStatistics
from backend.app.logging_config import get_logger

logger = get_logger("app.domains.rendering.engine")


ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocumentRenderer:
    def __init__(self) -> None:
        self._statistics: RenderingStatistics = RenderingStatistics()
        self._errors: list[tuple[RenderErrorCode, str]] = []

    @property
    def statistics(self) -> RenderingStatistics:
        return self._statistics

    @property
    def errors(self) -> list[tuple[RenderErrorCode, str]]:
        return self._errors

    def render(
        self, assembled_structure: dict[str, Any]
    ) -> tuple[bytes | None, RenderingStatistics]:
        self._statistics = RenderingStatistics()
        self._errors = []

        try:
            doc = Document()

            metadata = assembled_structure.get("metadata", {})
            self._apply_document_metadata(doc, metadata)

            blocks = assembled_structure.get("blocks", [])
            for block_data in blocks:
                self._render_block(doc, block_data)

            headers = assembled_structure.get("headers", [])
            for header_data in headers:
                self._render_header(doc, header_data)

            footers = assembled_structure.get("footers", [])
            for footer_data in footers:
                self._render_footer(doc, footer_data)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            content = buffer.read()

            self._statistics.total_blocks = (
                self._statistics.paragraphs
                + self._statistics.headings
                + self._statistics.tables
                + self._statistics.lists
                + self._statistics.page_breaks
                + self._statistics.section_breaks
            )

            return content, self._statistics

        except Exception as e:
            logger.error(f"Rendering failed: {e}")
            self._errors.append((RenderErrorCode.RENDERING_FAILED, str(e)))
            return None, self._statistics

    def _apply_document_metadata(self, doc: Document, metadata: dict[str, Any]) -> None:
        if not metadata:
            return

        core_props = doc.core_properties

        if metadata.get("title"):
            core_props.title = metadata["title"]
        if metadata.get("author"):
            core_props.author = metadata["author"]
        if metadata.get("subject"):
            core_props.subject = metadata["subject"]
        if metadata.get("keywords"):
            core_props.keywords = metadata["keywords"]

    def _render_block(self, doc: Document, block_data: dict[str, Any]) -> None:
        block_type = block_data.get("block_type", "")

        try:
            if block_type == "paragraph":
                self._render_paragraph(doc, block_data)
            elif block_type == "heading":
                self._render_heading(doc, block_data)
            elif block_type == "table":
                self._render_table(doc, block_data)
            elif block_type == "list":
                self._render_list(doc, block_data)
            elif block_type == "page_break":
                self._render_page_break(doc, block_data)
            elif block_type == "section_break":
                self._render_section_break(doc, block_data)
        except Exception as e:
            logger.error(f"Failed to render block {block_data.get('block_id')}: {e}")
            self._errors.append((RenderErrorCode.BLOCK_RENDERING_FAILED, str(e)))

    def _render_paragraph(self, doc: Document, block_data: dict[str, Any]) -> None:
        paragraph = doc.add_paragraph()

        style_name = block_data.get("style_name")
        if style_name:
            try:
                paragraph.style = style_name
            except KeyError:
                pass

        alignment = block_data.get("alignment")
        if alignment and alignment in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[alignment]

        pf = paragraph.paragraph_format

        if block_data.get("indent_left") is not None:
            pf.left_indent = Twips(int(block_data["indent_left"]))
        if block_data.get("indent_right") is not None:
            pf.right_indent = Twips(int(block_data["indent_right"]))
        if block_data.get("indent_first_line") is not None:
            pf.first_line_indent = Twips(int(block_data["indent_first_line"]))
        if block_data.get("spacing_before") is not None:
            pf.space_before = Twips(int(block_data["spacing_before"]))
        if block_data.get("spacing_after") is not None:
            pf.space_after = Twips(int(block_data["spacing_after"]))

        runs = block_data.get("runs", [])
        for run_data in runs:
            self._add_run(paragraph, run_data)

        self._statistics.paragraphs += 1

    def _render_heading(self, doc: Document, block_data: dict[str, Any]) -> None:
        level = block_data.get("level", 1)
        level = max(0, min(level, 9))

        paragraph = doc.add_heading(level=level)
        paragraph.clear()

        alignment = block_data.get("alignment")
        if alignment and alignment in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[alignment]

        runs = block_data.get("runs", [])
        for run_data in runs:
            self._add_run(paragraph, run_data)

        self._statistics.headings += 1

    def _render_table(self, doc: Document, block_data: dict[str, Any]) -> None:
        try:
            rows_data = block_data.get("rows", [])
            if not rows_data:
                return

            column_count = block_data.get("column_count", 0)
            if column_count == 0 and rows_data:
                first_row = rows_data[0]
                cells = first_row.get("cells", [])
                column_count = len(cells)

            if column_count == 0:
                return

            row_count = len(rows_data)
            table = doc.add_table(rows=row_count, cols=column_count)

            style_name = block_data.get("style_name")
            if style_name:
                try:
                    table.style = style_name
                except KeyError:
                    table.style = "Table Grid"
            else:
                table.style = "Table Grid"

            for row_idx, row_data in enumerate(rows_data):
                cells = row_data.get("cells", [])
                for cell_idx, cell_data in enumerate(cells):
                    if cell_idx >= column_count:
                        break

                    cell = table.cell(row_idx, cell_idx)
                    cell_content = cell_data.get("content", [])

                    cell.text = ""

                    for content_block in cell_content:
                        if content_block.get("block_type") == "paragraph":
                            cell_para = (
                                cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                            )
                            runs = content_block.get("runs", [])
                            for run_data in runs:
                                self._add_run(cell_para, run_data)

            self._statistics.tables += 1

        except Exception as e:
            logger.error(f"Failed to render table: {e}")
            self._errors.append((RenderErrorCode.TABLE_RENDERING_FAILED, str(e)))

    def _render_list(self, doc: Document, block_data: dict[str, Any]) -> None:
        try:
            list_type = block_data.get("list_type", "bullet")
            items = block_data.get("items", [])

            for item_data in items:
                if list_type == "bullet":
                    paragraph = doc.add_paragraph(style="List Bullet")
                else:
                    paragraph = doc.add_paragraph(style="List Number")

                paragraph.clear()

                runs = item_data.get("runs", [])
                for run_data in runs:
                    self._add_run(paragraph, run_data)

            self._statistics.lists += 1

        except Exception as e:
            logger.error(f"Failed to render list: {e}")
            self._errors.append((RenderErrorCode.LIST_RENDERING_FAILED, str(e)))

    def _render_page_break(self, doc: Document, block_data: dict[str, Any]) -> None:
        doc.add_page_break()
        self._statistics.page_breaks += 1

    def _render_section_break(self, doc: Document, block_data: dict[str, Any]) -> None:
        section = doc.add_section()
        orientation = block_data.get("orientation", "")
        page_width = block_data.get("page_width")
        page_height = block_data.get("page_height")

        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            if page_width and page_height:
                section.page_width = Twips(int(page_height))
                section.page_height = Twips(int(page_width))
        elif orientation == "portrait":
            section.orientation = WD_ORIENT.PORTRAIT
            if page_width and page_height:
                section.page_width = Twips(int(page_width))
                section.page_height = Twips(int(page_height))

        self._statistics.section_breaks += 1

    def _render_header(self, doc: Document, header_data: dict[str, Any]) -> None:
        try:
            section = doc.sections[-1] if doc.sections else doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False

            content = header_data.get("content", [])
            for block in content:
                if block.get("block_type") == "paragraph":
                    paragraph = header.add_paragraph()
                    runs = block.get("runs", [])
                    for run_data in runs:
                        self._add_run(paragraph, run_data)

            self._statistics.headers += 1

        except Exception as e:
            logger.error(f"Failed to render header: {e}")
            self._errors.append((RenderErrorCode.HEADER_FOOTER_RENDERING_FAILED, str(e)))

    def _render_footer(self, doc: Document, footer_data: dict[str, Any]) -> None:
        try:
            section = doc.sections[-1] if doc.sections else doc.sections[0]
            footer = section.footer
            footer.is_linked_to_previous = False

            content = footer_data.get("content", [])
            for block in content:
                if block.get("block_type") == "paragraph":
                    paragraph = footer.add_paragraph()
                    runs = block.get("runs", [])
                    for run_data in runs:
                        self._add_run(paragraph, run_data)

            self._statistics.footers += 1

        except Exception as e:
            logger.error(f"Failed to render footer: {e}")
            self._errors.append((RenderErrorCode.HEADER_FOOTER_RENDERING_FAILED, str(e)))

    def _add_run(self, paragraph, run_data: dict[str, Any]) -> None:
        text = run_data.get("text", "")
        run = paragraph.add_run(text)

        if run_data.get("bold"):
            run.bold = True
        if run_data.get("italic"):
            run.italic = True
        if run_data.get("underline"):
            run.underline = True
        if run_data.get("strike"):
            run.font.strike = True

        font_name = run_data.get("font_name")
        if font_name:
            run.font.name = font_name

        font_size = run_data.get("font_size")
        if font_size:
            run.font.size = Pt(font_size)

        color = run_data.get("color")
        if color:
            try:
                if color.startswith("#"):
                    color = color[1:]
                r = int(color[0:2], 16)
                g = int(color[2:4], 16)
                b = int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass

    def _apply_style(self, paragraph, style_name: str) -> bool:
        if not style_name:
            return False
        try:
            paragraph.style = style_name
            return True
        except KeyError:
            return False
