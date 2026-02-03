import io
import zipfile
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from backend.app.domains.rendering.schemas import RenderErrorCode, RenderingValidationResult
from backend.app.logging_config import get_logger

logger = get_logger("app.domains.rendering.validator")


class RenderedDocumentValidator:
    def validate(
        self,
        content: bytes,
        expected_blocks: int | None = None,
        expected_paragraphs: int | None = None,
        expected_tables: int | None = None,
        expected_headings: int | None = None,
    ) -> RenderingValidationResult:
        result = RenderingValidationResult()

        if not content:
            result.add_error(RenderErrorCode.VALIDATION_FAILED, "No content to validate")
            return result

        result.file_size_bytes = len(content)
        result.file_exists = True

        if not self._is_valid_zip(content):
            result.add_error(
                RenderErrorCode.FILE_CORRUPTION_DETECTED, "File is not a valid ZIP archive"
            )
            return result

        if not self._has_required_docx_parts(content):
            result.add_error(
                RenderErrorCode.FILE_CORRUPTION_DETECTED, "File is missing required DOCX components"
            )
            return result

        result.no_corruption = True

        try:
            doc = Document(io.BytesIO(content))
            result.file_opens_cleanly = True
        except PackageNotFoundError as e:
            result.add_error(RenderErrorCode.FILE_CORRUPTION_DETECTED, f"Cannot open document: {e}")
            return result
        except Exception as e:
            result.add_error(RenderErrorCode.FILE_CORRUPTION_DETECTED, f"Document open failed: {e}")
            return result

        try:
            stats = self._extract_document_statistics(doc)
            result.block_count = stats["total_blocks"]
            result.paragraph_count = stats["paragraphs"]
            result.table_count = stats["tables"]
            result.heading_count = stats["headings"]

            if result.block_count > 0:
                result.content_present = True

            result.structure_intact = True

            if expected_blocks is not None and result.block_count < expected_blocks:
                result.add_warning(
                    f"Expected {expected_blocks} blocks but found {result.block_count}"
                )

            if expected_paragraphs is not None and result.paragraph_count < expected_paragraphs:
                result.add_warning(
                    f"Expected {expected_paragraphs} paragraphs but found {result.paragraph_count}"
                )

            if expected_tables is not None and result.table_count != expected_tables:
                result.add_warning(
                    f"Expected {expected_tables} tables but found {result.table_count}"
                )

        except Exception as e:
            result.add_error(
                RenderErrorCode.VALIDATION_FAILED, f"Statistics extraction failed: {e}"
            )
            return result

        return result

    def validate_determinism(
        self,
        content1: bytes,
        content2: bytes,
    ) -> tuple[bool, str]:
        if content1 == content2:
            return True, "Content is byte-identical"

        try:
            doc1 = Document(io.BytesIO(content1))
            doc2 = Document(io.BytesIO(content2))

            paras1 = [p.text for p in doc1.paragraphs]
            paras2 = [p.text for p in doc2.paragraphs]

            if paras1 != paras2:
                return False, "Paragraph content differs between renders"

            tables1 = len(doc1.tables)
            tables2 = len(doc2.tables)

            if tables1 != tables2:
                return False, f"Table count differs: {tables1} vs {tables2}"

            for i, (t1, t2) in enumerate(zip(doc1.tables, doc2.tables)):
                if len(t1.rows) != len(t2.rows):
                    return False, f"Table {i} row count differs"
                for j, (r1, r2) in enumerate(zip(t1.rows, t2.rows)):
                    for k, (c1, c2) in enumerate(zip(r1.cells, r2.cells)):
                        if c1.text != c2.text:
                            return False, f"Table {i} cell ({j},{k}) content differs"

            return True, "Content is structurally identical"

        except Exception as e:
            return False, f"Comparison failed: {e}"

    def _is_valid_zip(self, content: bytes) -> bool:
        try:
            with zipfile.ZipFile(io.BytesIO(content), "r") as zf:
                return zf.testzip() is None
        except zipfile.BadZipFile:
            return False
        except Exception:
            return False

    def _has_required_docx_parts(self, content: bytes) -> bool:
        required_parts = [
            "[Content_Types].xml",
            "word/document.xml",
        ]

        try:
            with zipfile.ZipFile(io.BytesIO(content), "r") as zf:
                names = zf.namelist()
                for part in required_parts:
                    if part not in names:
                        return False
                return True
        except Exception:
            return False

    def _extract_document_statistics(self, doc: Document) -> dict[str, int]:
        paragraphs = len(doc.paragraphs)
        tables = len(doc.tables)

        headings = 0
        for para in doc.paragraphs:
            if para.style and para.style.name and para.style.name.startswith("Heading"):
                headings += 1

        total_blocks = paragraphs + tables

        return {
            "total_blocks": total_blocks,
            "paragraphs": paragraphs,
            "tables": tables,
            "headings": headings,
        }

    def validate_content_integrity(
        self,
        content: bytes,
        assembled_structure: dict[str, Any],
    ) -> RenderingValidationResult:
        result = self.validate(content)

        if not result.is_valid:
            return result

        expected_blocks = assembled_structure.get("blocks", [])

        doc = Document(io.BytesIO(content))

        assembled_texts = []
        for block in expected_blocks:
            runs = block.get("runs", [])
            block_text = "".join(run.get("text", "") for run in runs)
            if block_text.strip():
                assembled_texts.append(block_text.strip())

        rendered_texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                rendered_texts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        rendered_texts.append(cell.text.strip())

        missing_content = []
        for expected in assembled_texts:
            found = False
            for rendered in rendered_texts:
                if expected in rendered or rendered in expected:
                    found = True
                    break
            if not found:
                missing_content.append(expected[:50])

        if missing_content:
            result.add_warning(f"Some content may not have rendered: {missing_content[:3]}")

        return result
