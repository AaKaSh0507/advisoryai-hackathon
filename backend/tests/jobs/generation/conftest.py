from datetime import datetime
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock
from uuid import UUID, uuid4

import pytest
import pytest_asyncio
from docx import Document

from backend.app.domains.assembly.models import AssembledDocument, AssemblyStatus
from backend.app.domains.assembly.repository import AssembledDocumentRepository
from backend.app.domains.assembly.service import DocumentAssemblyService
from backend.app.domains.audit.generation_audit_service import GenerationAuditService
from backend.app.domains.audit.repository import AuditRepository
from backend.app.domains.document.models import Document as DocumentModel
from backend.app.domains.document.repository import DocumentRepository
from backend.app.domains.generation.models import (
    GenerationInput,
    GenerationInputBatch,
    GenerationInputStatus,
)
from backend.app.domains.generation.repository import GenerationInputRepository
from backend.app.domains.generation.section_output_models import (
    SectionGenerationStatus,
    SectionOutput,
    SectionOutputBatch,
)
from backend.app.domains.generation.section_output_repository import SectionOutputRepository
from backend.app.domains.generation.section_output_service import SectionGenerationService
from backend.app.domains.generation.service import GenerationInputService
from backend.app.domains.job.models import Job, JobStatus, JobType
from backend.app.domains.job.repository import JobRepository
from backend.app.domains.job.service import JobService
from backend.app.domains.rendering.models import RenderedDocument, RenderStatus
from backend.app.domains.rendering.repository import RenderedDocumentRepository
from backend.app.domains.rendering.service import DocumentRenderingService
from backend.app.domains.section.models import Section, SectionType
from backend.app.domains.section.repository import SectionRepository
from backend.app.domains.template.models import Template, TemplateVersion
from backend.app.domains.template.repository import TemplateRepository
from backend.app.domains.versioning.repository import VersioningRepository
from backend.app.domains.versioning.service import DocumentVersioningService
from backend.app.worker.handlers.base import HandlerContext
from backend.app.worker.handlers.generation_pipeline import GenerationPipelineHandler


@pytest.fixture
def fixed_document_id() -> UUID:
    return UUID("33333333-3333-3333-3333-333333333333")


@pytest.fixture
def fixed_template_id() -> UUID:
    return UUID("11111111-1111-1111-1111-111111111111")


@pytest.fixture
def fixed_template_version_id() -> UUID:
    return UUID("22222222-2222-2222-2222-222222222222")


@pytest.fixture
def fixed_job_id() -> UUID:
    return UUID("44444444-4444-4444-4444-444444444444")


@pytest.fixture
def fixed_input_batch_id() -> UUID:
    return UUID("55555555-5555-5555-5555-555555555555")


@pytest.fixture
def fixed_output_batch_id() -> UUID:
    return UUID("66666666-6666-6666-6666-666666666666")


@pytest.fixture
def fixed_assembled_document_id() -> UUID:
    return UUID("77777777-7777-7777-7777-777777777777")


@pytest.fixture
def fixed_rendered_document_id() -> UUID:
    return UUID("88888888-8888-8888-8888-888888888888")


@pytest.fixture
def fixed_version_id() -> UUID:
    return UUID("99999999-9999-9999-9999-999999999999")


@pytest.fixture
def sample_client_data() -> dict:
    return {
        "client_id": "client_001",
        "client_name": "Test Corporation",
        "data_fields": {
            "project_name": "Alpha Project",
            "contract_value": 1500000,
            "start_date": "2026-01-15",
        },
        "custom_context": {
            "industry": "Technology",
            "region": "North America",
        },
    }


@pytest.fixture
def sample_job_payload(
    fixed_template_version_id: UUID,
    fixed_document_id: UUID,
    sample_client_data: dict,
) -> dict:
    return {
        "template_version_id": str(fixed_template_version_id),
        "document_id": str(fixed_document_id),
        "version_intent": 1,
        "client_data": sample_client_data,
        "force_regenerate": False,
    }


@pytest.fixture
def sample_job(fixed_job_id: UUID, sample_job_payload: dict) -> Job:
    job = Job(
        id=fixed_job_id,
        job_type=JobType.GENERATE,
        payload=sample_job_payload,
        status=JobStatus.RUNNING,
        worker_id="worker-test-1",
    )
    job.created_at = datetime(2026, 1, 15, 10, 0, 0)
    job.started_at = datetime(2026, 1, 15, 10, 0, 1)
    return job


@pytest.fixture
def sample_generation_input_batch(
    fixed_input_batch_id: UUID,
    fixed_document_id: UUID,
    fixed_template_version_id: UUID,
) -> GenerationInputBatch:
    batch = GenerationInputBatch(
        id=fixed_input_batch_id,
        document_id=fixed_document_id,
        template_version_id=fixed_template_version_id,
        version_intent=1,
        status=GenerationInputStatus.VALIDATED,
        content_hash="abc123def456",
        total_inputs=3,
        is_immutable=True,
    )
    batch.created_at = datetime(2026, 1, 15, 10, 0, 0)
    batch.validated_at = datetime(2026, 1, 15, 10, 0, 1)
    return batch


@pytest.fixture
def sample_generation_inputs(
    fixed_input_batch_id: UUID,
    fixed_template_version_id: UUID,
) -> list[GenerationInput]:
    inputs = []
    section_configs = [
        (1, "body/introduction", 0),
        (2, "body/main_content", 1),
        (3, "body/conclusion", 2),
    ]
    for section_id, path, seq in section_configs:
        inp = GenerationInput(
            id=uuid4(),
            batch_id=fixed_input_batch_id,
            section_id=section_id,
            sequence_order=seq,
            template_id=fixed_template_version_id,
            template_version_id=fixed_template_version_id,
            structural_path=path,
            hierarchy_context={
                "parent_heading": "body",
                "depth": 1,
                "path_segments": ["body", path.split("/")[-1]],
            },
            prompt_config={
                "classification_confidence": 0.95,
                "classification_method": "RULE_BASED",
                "justification": f"Dynamic content for {path}",
            },
            client_data={
                "client_name": "Test Corporation",
                "data_fields": {"project": "Alpha Project"},
            },
            surrounding_context={},
            input_hash=f"hash_{section_id}",
        )
        inp.created_at = datetime(2026, 1, 15, 10, 0, 0)
        inputs.append(inp)
    return inputs


@pytest.fixture
def sample_section_output_batch(
    fixed_output_batch_id: UUID,
    fixed_input_batch_id: UUID,
) -> SectionOutputBatch:
    batch = SectionOutputBatch(
        id=fixed_output_batch_id,
        input_batch_id=fixed_input_batch_id,
        status=SectionGenerationStatus.COMPLETED,
        total_sections=3,
        completed_count=3,
        failed_count=0,
    )
    batch.created_at = datetime(2026, 1, 15, 10, 0, 2)
    batch.completed_at = datetime(2026, 1, 15, 10, 0, 5)
    return batch


@pytest.fixture
def sample_section_outputs(
    fixed_output_batch_id: UUID,
    sample_generation_inputs: list[GenerationInput],
) -> list[SectionOutput]:
    outputs = []
    for inp in sample_generation_inputs:
        output = SectionOutput(
            id=uuid4(),
            batch_id=fixed_output_batch_id,
            input_id=inp.id,
            section_id=inp.section_id,
            status=SectionGenerationStatus.COMPLETED,
            generated_content=f"Generated content for section {inp.section_id}",
            token_count=50,
            retry_count=0,
        )
        output.created_at = datetime(2026, 1, 15, 10, 0, 3)
        output.completed_at = datetime(2026, 1, 15, 10, 0, 4)
        outputs.append(output)
    return outputs


@pytest.fixture
def sample_assembled_document(
    fixed_assembled_document_id: UUID,
    fixed_document_id: UUID,
    fixed_template_version_id: UUID,
    fixed_output_batch_id: UUID,
) -> AssembledDocument:
    doc = AssembledDocument(
        id=fixed_assembled_document_id,
        document_id=fixed_document_id,
        template_version_id=fixed_template_version_id,
        version_intent=1,
        section_output_batch_id=fixed_output_batch_id,
        status=AssemblyStatus.COMPLETED,
        total_blocks=5,
        static_blocks_count=2,
        dynamic_blocks_count=3,
        assembly_hash="assembled_hash_123",
    )
    doc.created_at = datetime(2026, 1, 15, 10, 0, 6)
    doc.assembled_at = datetime(2026, 1, 15, 10, 0, 7)
    return doc


@pytest.fixture
def sample_rendered_document(
    fixed_rendered_document_id: UUID,
    fixed_assembled_document_id: UUID,
    fixed_document_id: UUID,
) -> RenderedDocument:
    doc = RenderedDocument(
        id=fixed_rendered_document_id,
        assembled_document_id=fixed_assembled_document_id,
        document_id=fixed_document_id,
        version=1,
        status=RenderStatus.COMPLETED,
        output_path="documents/33333333-3333-3333-3333-333333333333/1/output.docx",
        file_size_bytes=12345,
        content_hash="rendered_hash_456",
    )
    doc.created_at = datetime(2026, 1, 15, 10, 0, 8)
    doc.rendered_at = datetime(2026, 1, 15, 10, 0, 9)
    return doc


@pytest.fixture
def sample_docx_content():
    def _create(text: str = "Sample generated document content"):
        doc = Document()
        doc.add_heading("Generated Document", level=1)
        doc.add_paragraph(text)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    return _create


@pytest.fixture
def pipeline_handler() -> GenerationPipelineHandler:
    return GenerationPipelineHandler()


@pytest.fixture
def mock_generation_input_service() -> MagicMock:
    mock_service = MagicMock(spec=GenerationInputService)
    mock_service.prepare_generation_inputs = AsyncMock()
    return mock_service


@pytest.fixture
def mock_section_generation_service() -> MagicMock:
    mock_service = MagicMock(spec=SectionGenerationService)
    mock_service.execute_section_generation = AsyncMock()
    return mock_service


@pytest.fixture
def mock_assembly_service() -> MagicMock:
    mock_service = MagicMock(spec=DocumentAssemblyService)
    mock_service.assemble_document = AsyncMock()
    return mock_service


@pytest.fixture
def mock_rendering_service() -> MagicMock:
    mock_service = MagicMock(spec=DocumentRenderingService)
    mock_service.render_document = AsyncMock()
    return mock_service


@pytest.fixture
def mock_versioning_service() -> MagicMock:
    mock_service = MagicMock(spec=DocumentVersioningService)
    mock_service.create_version = AsyncMock()
    return mock_service


@pytest.fixture
def mock_generation_audit_service() -> MagicMock:
    mock_service = MagicMock(spec=GenerationAuditService)
    mock_service.log_pipeline_started = AsyncMock()
    mock_service.log_stage_completed = AsyncMock()
    mock_service.log_pipeline_completed = AsyncMock()
    mock_service.log_pipeline_failed = AsyncMock()
    return mock_service


@pytest.fixture
def mock_storage_service(sample_docx_content) -> MagicMock:
    mock_storage = MagicMock()
    mock_storage.get_file = MagicMock(return_value=sample_docx_content())
    mock_storage.upload_file = MagicMock(return_value="documents/test/output.docx")
    mock_storage.file_exists = MagicMock(return_value=True)
    return mock_storage


class MockPrepareGenerationInputsResponse:
    def __init__(self, batch_id: UUID, total_dynamic_sections: int = 3):
        self.batch_id = batch_id
        self.total_dynamic_sections = total_dynamic_sections
        self.validation_status = "VALIDATED"


class MockSectionGenerationResponse:
    def __init__(
        self,
        batch_id: UUID,
        total_sections: int = 3,
        completed_count: int = 3,
        failed_count: int = 0,
    ):
        self.batch_id = batch_id
        self.total_sections = total_sections
        self.completed_count = completed_count
        self.failed_count = failed_count


class MockAssemblyResult:
    def __init__(
        self,
        success: bool = True,
        assembled_document: AssembledDocument | None = None,
        error_message: str | None = None,
    ):
        self.success = success
        self.assembled_document = assembled_document
        self.error_message = error_message


class MockRenderingResult:
    def __init__(
        self,
        success: bool = True,
        rendered_document: RenderedDocument | None = None,
        output_path: str | None = None,
        error_message: str | None = None,
    ):
        self.success = success
        self.rendered_document = rendered_document
        self.output_path = output_path
        self.error_message = error_message


class MockVersionResult:
    def __init__(
        self,
        success: bool = True,
        version_id: UUID | None = None,
        version_number: int | None = None,
        output_path: str | None = None,
        error: MagicMock | None = None,
    ):
        self.success = success
        self.version_id = version_id
        self.version_number = version_number
        self.output_path = output_path
        self.error = error


@pytest.fixture
def mock_prepare_response(fixed_input_batch_id: UUID) -> MockPrepareGenerationInputsResponse:
    return MockPrepareGenerationInputsResponse(batch_id=fixed_input_batch_id)


@pytest.fixture
def mock_generation_response(fixed_output_batch_id: UUID) -> MockSectionGenerationResponse:
    return MockSectionGenerationResponse(batch_id=fixed_output_batch_id)


@pytest.fixture
def mock_assembly_result(sample_assembled_document: AssembledDocument) -> MockAssemblyResult:
    return MockAssemblyResult(success=True, assembled_document=sample_assembled_document)


@pytest.fixture
def mock_rendering_result(sample_rendered_document: RenderedDocument) -> MockRenderingResult:
    return MockRenderingResult(
        success=True,
        rendered_document=sample_rendered_document,
        output_path=sample_rendered_document.output_path,
    )


@pytest.fixture
def mock_version_result(
    fixed_version_id: UUID,
) -> MockVersionResult:
    return MockVersionResult(
        success=True,
        version_id=fixed_version_id,
        version_number=1,
        output_path="documents/33333333-3333-3333-3333-333333333333/versions/1/final.docx",
    )


@pytest_asyncio.fixture
async def db_job_repository(db_session) -> JobRepository:
    return JobRepository(db_session)


@pytest_asyncio.fixture
async def db_job_service(db_job_repository: JobRepository) -> JobService:
    return JobService(db_job_repository)


@pytest_asyncio.fixture
async def db_template_repository(db_session) -> TemplateRepository:
    return TemplateRepository(db_session)


@pytest_asyncio.fixture
async def db_document_repository(db_session) -> DocumentRepository:
    return DocumentRepository(db_session)


@pytest_asyncio.fixture
async def db_section_repository(db_session) -> SectionRepository:
    return SectionRepository(db_session)


@pytest_asyncio.fixture
async def db_generation_input_repository(db_session) -> GenerationInputRepository:
    return GenerationInputRepository(db_session)


@pytest_asyncio.fixture
async def db_section_output_repository(db_session) -> SectionOutputRepository:
    return SectionOutputRepository(db_session)


@pytest_asyncio.fixture
async def db_assembled_document_repository(db_session) -> AssembledDocumentRepository:
    return AssembledDocumentRepository(db_session)


@pytest_asyncio.fixture
async def db_rendered_document_repository(db_session) -> RenderedDocumentRepository:
    return RenderedDocumentRepository(db_session)


@pytest_asyncio.fixture
async def db_versioning_repository(db_session) -> VersioningRepository:
    return VersioningRepository(db_session)


@pytest_asyncio.fixture
async def db_audit_repository(db_session) -> AuditRepository:
    return AuditRepository(db_session)


@pytest_asyncio.fixture
async def setup_template_and_document(
    db_session,
    db_template_repository: TemplateRepository,
    db_document_repository: DocumentRepository,
    db_section_repository: SectionRepository,
    fixed_template_id: UUID,
    fixed_template_version_id: UUID,
    fixed_document_id: UUID,
):
    template = Template(id=fixed_template_id, name="Test Template for Pipeline")
    await db_template_repository.create(template)

    version = TemplateVersion(
        id=fixed_template_version_id,
        template_id=fixed_template_id,
        version_number=1,
        source_doc_path=f"templates/{fixed_template_id}/1/source.docx",
        parsed_doc_path=f"templates/{fixed_template_id}/1/parsed.json",
    )
    await db_template_repository.create_version(version)

    document = DocumentModel(
        id=fixed_document_id,
        template_id=fixed_template_id,
        name="Test Document for Pipeline",
    )
    await db_document_repository.create(document)

    sections = [
        Section(
            id=1,
            template_version_id=fixed_template_version_id,
            block_id="blk_hea_001",
            structural_path="header/title",
            section_type=SectionType.STATIC,
            sequence_order=0,
            original_content={"block_type": "heading", "text": "Document Title"},
        ),
        Section(
            id=2,
            template_version_id=fixed_template_version_id,
            block_id="blk_par_001",
            structural_path="body/introduction",
            section_type=SectionType.DYNAMIC,
            sequence_order=1,
            original_content={"block_type": "paragraph", "text": "[PLACEHOLDER]"},
        ),
        Section(
            id=3,
            template_version_id=fixed_template_version_id,
            block_id="blk_par_002",
            structural_path="body/main_content",
            section_type=SectionType.DYNAMIC,
            sequence_order=2,
            original_content={"block_type": "paragraph", "text": "[PLACEHOLDER]"},
        ),
        Section(
            id=4,
            template_version_id=fixed_template_version_id,
            block_id="blk_par_003",
            structural_path="body/conclusion",
            section_type=SectionType.DYNAMIC,
            sequence_order=3,
            original_content={"block_type": "paragraph", "text": "[PLACEHOLDER]"},
        ),
        Section(
            id=5,
            template_version_id=fixed_template_version_id,
            block_id="blk_par_004",
            structural_path="footer/disclaimer",
            section_type=SectionType.STATIC,
            sequence_order=4,
            original_content={"block_type": "paragraph", "text": "Standard disclaimer text"},
        ),
    ]
    for section in sections:
        await db_section_repository.create(section)

    await db_session.commit()

    return {
        "template": template,
        "template_version": version,
        "document": document,
        "sections": sections,
    }


def create_handler_context(
    session,
    job: Job,
) -> HandlerContext:
    return HandlerContext(session=session, job=job)
