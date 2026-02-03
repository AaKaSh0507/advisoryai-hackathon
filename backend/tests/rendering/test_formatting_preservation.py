import io

from docx import Document

from backend.app.domains.rendering.engine import DocumentRenderer


class TestStaticFormattingPreserved:
    def test_bold_text_is_preserved(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [
                        {"text": "Bold text", "bold": True, "italic": False},
                    ],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        found_bold = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "Bold text" in run.text and run.bold:
                    found_bold = True
                    break

        assert found_bold

    def test_italic_text_is_preserved(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [
                        {"text": "Italic text", "bold": False, "italic": True},
                    ],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        found_italic = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "Italic text" in run.text and run.italic:
                    found_italic = True
                    break

        assert found_italic

    def test_underline_text_is_preserved(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [
                        {"text": "Underline text", "underline": True},
                    ],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        found_underline = False
        for para in doc.paragraphs:
            for run in para.runs:
                if "Underline text" in run.text and run.underline:
                    found_underline = True
                    break

        assert found_underline

    def test_multiple_formatting_in_same_paragraph(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [
                        {"text": "Normal ", "bold": False, "italic": False},
                        {"text": "bold ", "bold": True, "italic": False},
                        {"text": "italic", "bold": False, "italic": True},
                    ],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        para = doc.paragraphs[0]
        assert len(para.runs) == 3

        run_texts = [r.text for r in para.runs]
        assert "Normal " in run_texts
        assert "bold " in run_texts
        assert "italic" in run_texts


class TestDynamicContentDoesNotAlterStyles:
    def test_injected_content_uses_paragraph_style(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [
                        {"text": "Injected dynamic content"},
                    ],
                    "style_name": "Normal",
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        para = doc.paragraphs[0]
        assert para.style is not None

    def test_static_block_formatting_unchanged_after_dynamic_injection(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_static",
                    "sequence": 0,
                    "runs": [
                        {"text": "Static bold text", "bold": True},
                    ],
                },
                {
                    "block_type": "paragraph",
                    "block_id": "blk_dynamic",
                    "sequence": 1,
                    "runs": [
                        {"text": "Dynamic content"},
                    ],
                },
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        first_para = doc.paragraphs[0]
        for run in first_para.runs:
            if "Static bold text" in run.text:
                assert run.bold


class TestTablesRemainIntact:
    def test_table_row_count_preserved(
        self,
        document_renderer: DocumentRenderer,
        sample_table_block: dict,
    ):
        structure = {
            "blocks": [sample_table_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2

    def test_table_column_count_preserved(
        self,
        document_renderer: DocumentRenderer,
        sample_table_block: dict,
    ):
        structure = {
            "blocks": [sample_table_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        table = doc.tables[0]
        for row in table.rows:
            assert len(row.cells) == 2

    def test_table_cell_content_preserved(
        self,
        document_renderer: DocumentRenderer,
        sample_table_block: dict,
    ):
        structure = {
            "blocks": [sample_table_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        table = doc.tables[0]
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text.strip())

        assert "Header 1" in cell_texts
        assert "Header 2" in cell_texts
        assert "Value 1" in cell_texts
        assert "Value 2" in cell_texts


class TestNumberingRemainIntact:
    def test_bullet_list_renders_all_items(
        self,
        document_renderer: DocumentRenderer,
        sample_list_block: dict,
    ):
        structure = {
            "blocks": [sample_list_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "First bullet item" in all_text
        assert "Second bullet item" in all_text

    def test_numbered_list_renders_correctly(
        self,
        document_renderer: DocumentRenderer,
    ):
        numbered_list = {
            "block_type": "list",
            "block_id": "blk_lst_001",
            "sequence": 0,
            "list_type": "number",
            "items": [
                {"item_id": "item_001", "level": 0, "runs": [{"text": "First numbered item"}]},
                {"item_id": "item_002", "level": 0, "runs": [{"text": "Second numbered item"}]},
            ],
        }

        structure = {
            "blocks": [numbered_list],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, _ = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "First numbered item" in all_text
        assert "Second numbered item" in all_text


class TestHeadersRemainIntact:
    def test_header_content_is_rendered(
        self,
        document_renderer: DocumentRenderer,
        sample_header_data: dict,
    ):
        structure = {
            "blocks": [],
            "metadata": {},
            "headers": [sample_header_data],
            "footers": [],
        }

        content, stats = document_renderer.render(structure)

        assert stats.headers == 1

    def test_footer_content_is_rendered(
        self,
        document_renderer: DocumentRenderer,
        sample_footer_data: dict,
    ):
        structure = {
            "blocks": [],
            "metadata": {},
            "headers": [],
            "footers": [sample_footer_data],
        }

        content, stats = document_renderer.render(structure)

        assert stats.footers == 1


class TestHeadingLevelsPreserved:
    def test_heading_level_1_rendered(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "heading",
                    "block_id": "blk_hea_001",
                    "sequence": 0,
                    "level": 1,
                    "runs": [{"text": "Heading Level 1"}],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, stats = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        assert stats.headings == 1

        found = any("Heading Level 1" in p.text for p in doc.paragraphs)
        assert found

    def test_multiple_heading_levels_preserved(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure = {
            "blocks": [
                {
                    "block_type": "heading",
                    "block_id": "blk_hea_001",
                    "sequence": 0,
                    "level": 1,
                    "runs": [{"text": "Main Title"}],
                },
                {
                    "block_type": "heading",
                    "block_id": "blk_hea_002",
                    "sequence": 1,
                    "level": 2,
                    "runs": [{"text": "Sub Title"}],
                },
                {
                    "block_type": "heading",
                    "block_id": "blk_hea_003",
                    "sequence": 2,
                    "level": 3,
                    "runs": [{"text": "Section Title"}],
                },
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content, stats = document_renderer.render(structure)
        doc = Document(io.BytesIO(content))

        assert stats.headings == 3

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Main Title" in all_text
        assert "Sub Title" in all_text
        assert "Section Title" in all_text
