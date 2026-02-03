from datetime import datetime
from io import BytesIO
from unittest.mock import AsyncMock, MagicMock
from uuid import UUID, uuid4

import pytest
from docx import Document

from backend.app.domains.assembly.models import AssembledDocument, AssemblyStatus
from backend.app.domains.rendering.engine import DocumentRenderer
from backend.app.domains.rendering.models import RenderedDocument, RenderStatus
from backend.app.domains.rendering.repository import RenderedDocumentRepository
from backend.app.domains.rendering.schemas import RenderingRequest
from backend.app.domains.rendering.service import DocumentRenderingService
from backend.app.domains.rendering.validator import RenderedDocumentValidator


@pytest.fixture
def document_id() -> UUID:
    return uuid4()


@pytest.fixture
def assembled_document_id() -> UUID:
    return uuid4()


@pytest.fixture
def template_version_id() -> UUID:
    return uuid4()


@pytest.fixture
def section_output_batch_id() -> UUID:
    return uuid4()


@pytest.fixture
def sample_paragraph_block() -> dict:
    return {
        "block_type": "paragraph",
        "block_id": "blk_par_0001_abc123",
        "sequence": 0,
        "runs": [
            {
                "text": "This is a test paragraph with some content.",
                "bold": False,
                "italic": False,
                "underline": False,
                "strike": False,
            }
        ],
        "alignment": "left",
        "style_name": "Normal",
    }


@pytest.fixture
def sample_heading_block() -> dict:
    return {
        "block_type": "heading",
        "block_id": "blk_hea_0001_def456",
        "sequence": 1,
        "level": 1,
        "runs": [
            {
                "text": "Executive Summary",
                "bold": True,
                "italic": False,
            }
        ],
        "style_name": "Heading1",
    }


@pytest.fixture
def sample_table_block() -> dict:
    return {
        "block_type": "table",
        "block_id": "blk_tab_0001_ghi789",
        "sequence": 2,
        "column_count": 2,
        "rows": [
            {
                "row_id": "row_001",
                "row_index": 0,
                "is_header": True,
                "cells": [
                    {
                        "cell_id": "cell_001",
                        "row_index": 0,
                        "col_index": 0,
                        "content": [
                            {
                                "block_type": "paragraph",
                                "block_id": "blk_par_cell_001",
                                "sequence": 0,
                                "runs": [{"text": "Header 1"}],
                            }
                        ],
                    },
                    {
                        "cell_id": "cell_002",
                        "row_index": 0,
                        "col_index": 1,
                        "content": [
                            {
                                "block_type": "paragraph",
                                "block_id": "blk_par_cell_002",
                                "sequence": 0,
                                "runs": [{"text": "Header 2"}],
                            }
                        ],
                    },
                ],
            },
            {
                "row_id": "row_002",
                "row_index": 1,
                "is_header": False,
                "cells": [
                    {
                        "cell_id": "cell_003",
                        "row_index": 1,
                        "col_index": 0,
                        "content": [
                            {
                                "block_type": "paragraph",
                                "block_id": "blk_par_cell_003",
                                "sequence": 0,
                                "runs": [{"text": "Value 1"}],
                            }
                        ],
                    },
                    {
                        "cell_id": "cell_004",
                        "row_index": 1,
                        "col_index": 1,
                        "content": [
                            {
                                "block_type": "paragraph",
                                "block_id": "blk_par_cell_004",
                                "sequence": 0,
                                "runs": [{"text": "Value 2"}],
                            }
                        ],
                    },
                ],
            },
        ],
        "style_name": "Table Grid",
    }


@pytest.fixture
def sample_list_block() -> dict:
    return {
        "block_type": "list",
        "block_id": "blk_lst_0001_jkl012",
        "sequence": 3,
        "list_type": "bullet",
        "items": [
            {
                "item_id": "item_001",
                "level": 0,
                "runs": [{"text": "First bullet item"}],
                "bullet_char": "•",
            },
            {
                "item_id": "item_002",
                "level": 0,
                "runs": [{"text": "Second bullet item"}],
                "bullet_char": "•",
            },
        ],
        "style_name": "ListBullet",
    }


@pytest.fixture
def sample_header_data() -> dict:
    return {
        "block_type": "header",
        "block_id": "blk_hdr_0001",
        "sequence": 0,
        "header_footer_type": "default",
        "content": [
            {
                "block_type": "paragraph",
                "block_id": "blk_par_hdr_001",
                "sequence": 0,
                "runs": [{"text": "Document Header"}],
            }
        ],
    }


@pytest.fixture
def sample_footer_data() -> dict:
    return {
        "block_type": "footer",
        "block_id": "blk_ftr_0001",
        "sequence": 0,
        "header_footer_type": "default",
        "content": [
            {
                "block_type": "paragraph",
                "block_id": "blk_par_ftr_001",
                "sequence": 0,
                "runs": [{"text": "Page Footer"}],
            }
        ],
    }


@pytest.fixture
def sample_metadata() -> dict:
    return {
        "title": "Advisory Report",
        "author": "AdvisoryAI",
        "subject": "Financial Analysis",
        "keywords": "advisory, finance, report",
    }


@pytest.fixture
def simple_assembled_structure(
    sample_paragraph_block,
    sample_heading_block,
    sample_metadata,
) -> dict:
    return {
        "blocks": [sample_heading_block, sample_paragraph_block],
        "metadata": sample_metadata,
        "headers": [],
        "footers": [],
    }


@pytest.fixture
def complex_assembled_structure(
    sample_paragraph_block,
    sample_heading_block,
    sample_table_block,
    sample_list_block,
    sample_header_data,
    sample_footer_data,
    sample_metadata,
) -> dict:
    return {
        "blocks": [
            sample_heading_block,
            sample_paragraph_block,
            sample_table_block,
            sample_list_block,
        ],
        "metadata": sample_metadata,
        "headers": [sample_header_data],
        "footers": [sample_footer_data],
    }


@pytest.fixture
def rendering_request(
    assembled_document_id: UUID,
    document_id: UUID,
) -> RenderingRequest:
    return RenderingRequest(
        assembled_document_id=assembled_document_id,
        document_id=document_id,
        version=1,
    )


@pytest.fixture
def mock_assembled_document(
    assembled_document_id: UUID,
    document_id: UUID,
    template_version_id: UUID,
    section_output_batch_id: UUID,
    complex_assembled_structure: dict,
) -> MagicMock:
    mock = MagicMock(spec=AssembledDocument)
    mock.id = assembled_document_id
    mock.document_id = document_id
    mock.template_version_id = template_version_id
    mock.section_output_batch_id = section_output_batch_id
    mock.version_intent = 1
    mock.status = AssemblyStatus.VALIDATED
    mock.is_immutable = True
    mock.assembly_hash = "test_hash_123"
    mock.assembled_structure = {"blocks": complex_assembled_structure["blocks"]}
    mock.document_metadata = complex_assembled_structure["metadata"]
    mock.headers = complex_assembled_structure["headers"]
    mock.footers = complex_assembled_structure["footers"]
    mock.total_blocks = len(complex_assembled_structure["blocks"])
    return mock


@pytest.fixture
def mock_rendered_document(
    assembled_document_id: UUID,
    document_id: UUID,
) -> MagicMock:
    mock = MagicMock(spec=RenderedDocument)
    mock.id = uuid4()
    mock.assembled_document_id = assembled_document_id
    mock.document_id = document_id
    mock.version = 1
    mock.status = RenderStatus.PENDING
    mock.output_path = None
    mock.content_hash = None
    mock.file_size_bytes = 0
    mock.total_blocks_rendered = 0
    mock.paragraphs_rendered = 0
    mock.tables_rendered = 0
    mock.lists_rendered = 0
    mock.headings_rendered = 0
    mock.headers_rendered = 0
    mock.footers_rendered = 0
    mock.is_immutable = False
    mock.rendered_at = None
    mock.created_at = datetime.utcnow()
    return mock


@pytest.fixture
def mock_repository() -> AsyncMock:
    repo = AsyncMock(spec=RenderedDocumentRepository)
    repo.get_by_assembled_document.return_value = None
    repo.exists_for_assembled_document.return_value = False
    return repo


@pytest.fixture
def mock_assembled_repository(mock_assembled_document) -> AsyncMock:
    repo = AsyncMock()
    repo.get_by_id.return_value = mock_assembled_document
    return repo


@pytest.fixture
def document_renderer() -> DocumentRenderer:
    return DocumentRenderer()


@pytest.fixture
def document_validator() -> RenderedDocumentValidator:
    return RenderedDocumentValidator()


@pytest.fixture
def rendering_service(
    mock_repository: AsyncMock,
    mock_assembled_repository: AsyncMock,
    mock_storage,
) -> DocumentRenderingService:
    return DocumentRenderingService(
        repository=mock_repository,
        assembled_document_repository=mock_assembled_repository,
        storage=mock_storage,
    )


def create_valid_docx() -> bytes:
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test paragraph.")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def create_invalid_content() -> bytes:
    return b"This is not a valid docx file"
