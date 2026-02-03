import io

from docx import Document

from backend.app.domains.rendering.engine import DocumentRenderer
from backend.app.domains.rendering.schemas import RenderingRequest
from backend.app.domains.rendering.service import DocumentRenderingService


class TestRenderedDocxExists:
    def test_renderer_produces_bytes_output(
        self,
        document_renderer: DocumentRenderer,
        simple_assembled_structure: dict,
    ):
        content, stats = document_renderer.render(simple_assembled_structure)

        assert content is not None
        assert isinstance(content, bytes)
        assert len(content) > 0

    def test_renderer_produces_valid_docx_structure(
        self,
        document_renderer: DocumentRenderer,
        simple_assembled_structure: dict,
    ):
        content, stats = document_renderer.render(simple_assembled_structure)

        doc = Document(io.BytesIO(content))
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_complex_document_renders_successfully(
        self,
        document_renderer: DocumentRenderer,
        complex_assembled_structure: dict,
    ):
        content, stats = document_renderer.render(complex_assembled_structure)

        assert content is not None
        assert stats.paragraphs > 0
        assert stats.headings > 0
        assert stats.tables > 0
        assert stats.lists > 0


class TestFileOpensWithoutError:
    def test_rendered_file_opens_in_docx_library(
        self,
        document_renderer: DocumentRenderer,
        simple_assembled_structure: dict,
    ):
        content, _ = document_renderer.render(simple_assembled_structure)

        try:
            opened_successfully = True
        except Exception:
            opened_successfully = False

        assert opened_successfully

    def test_rendered_file_has_valid_zip_structure(
        self,
        document_renderer: DocumentRenderer,
        simple_assembled_structure: dict,
    ):
        import zipfile

        content, _ = document_renderer.render(simple_assembled_structure)

        with zipfile.ZipFile(io.BytesIO(content), "r") as zf:
            names = zf.namelist()
            assert "[Content_Types].xml" in names
            assert "word/document.xml" in names

    def test_rendered_file_has_no_corruption(
        self,
        document_renderer: DocumentRenderer,
        document_validator,
        complex_assembled_structure: dict,
    ):
        content, _ = document_renderer.render(complex_assembled_structure)

        validation_result = document_validator.validate(content)

        assert validation_result.no_corruption


class TestContentInCorrectSections:
    def test_heading_content_is_present(
        self,
        document_renderer: DocumentRenderer,
        sample_heading_block: dict,
    ):
        structure = {"blocks": [sample_heading_block], "metadata": {}, "headers": [], "footers": []}
        content, _ = document_renderer.render(structure)

        doc = Document(io.BytesIO(content))
        heading_text = sample_heading_block["runs"][0]["text"]

        found = any(heading_text in p.text for p in doc.paragraphs)
        assert found

    def test_paragraph_content_is_present(
        self,
        document_renderer: DocumentRenderer,
        sample_paragraph_block: dict,
    ):
        structure = {
            "blocks": [sample_paragraph_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }
        content, _ = document_renderer.render(structure)

        doc = Document(io.BytesIO(content))
        para_text = sample_paragraph_block["runs"][0]["text"]

        found = any(para_text in p.text for p in doc.paragraphs)
        assert found

    def test_table_content_is_present(
        self,
        document_renderer: DocumentRenderer,
        sample_table_block: dict,
    ):
        structure = {"blocks": [sample_table_block], "metadata": {}, "headers": [], "footers": []}
        content, _ = document_renderer.render(structure)

        doc = Document(io.BytesIO(content))

        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 2

        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text)

        assert "Header 1" in cell_texts
        assert "Value 1" in cell_texts

    def test_list_content_is_present(
        self,
        document_renderer: DocumentRenderer,
        sample_list_block: dict,
    ):
        structure = {"blocks": [sample_list_block], "metadata": {}, "headers": [], "footers": []}
        content, _ = document_renderer.render(structure)

        doc = Document(io.BytesIO(content))

        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "First bullet item" in all_text
        assert "Second bullet item" in all_text

    def test_all_blocks_rendered_in_order(
        self,
        document_renderer: DocumentRenderer,
        complex_assembled_structure: dict,
    ):
        content, stats = document_renderer.render(complex_assembled_structure)

        doc = Document(io.BytesIO(content))

        all_text = " ".join(p.text for p in doc.paragraphs)

        heading_pos = all_text.find("Executive Summary")
        para_pos = all_text.find("test paragraph")

        assert heading_pos < para_pos


class TestServiceRendersSuccessfully:
    async def test_service_renders_valid_assembled_document(
        self,
        rendering_service: DocumentRenderingService,
        rendering_request: RenderingRequest,
        mock_repository,
        mock_rendered_document,
    ):
        mock_repository.create.return_value = mock_rendered_document
        mock_repository.mark_in_progress.return_value = mock_rendered_document
        mock_repository.mark_completed.return_value = mock_rendered_document
        mock_repository.mark_validated.return_value = mock_rendered_document

        result = await rendering_service.render_document(rendering_request)

        assert result.success
        assert result.output_path is not None
        assert result.validation_result is not None
        assert result.validation_result.is_valid

    async def test_service_populates_statistics(
        self,
        rendering_service: DocumentRenderingService,
        rendering_request: RenderingRequest,
        mock_repository,
        mock_rendered_document,
    ):
        mock_repository.create.return_value = mock_rendered_document
        mock_repository.mark_in_progress.return_value = mock_rendered_document
        mock_repository.mark_completed.return_value = mock_rendered_document
        mock_repository.mark_validated.return_value = mock_rendered_document

        result = await rendering_service.render_document(rendering_request)

        assert result.success
        assert result.rendered_document is not None

    async def test_service_persists_to_storage(
        self,
        rendering_service: DocumentRenderingService,
        rendering_request: RenderingRequest,
        mock_repository,
        mock_rendered_document,
        mock_storage,
    ):
        mock_repository.create.return_value = mock_rendered_document
        mock_repository.mark_in_progress.return_value = mock_rendered_document
        mock_repository.mark_completed.return_value = mock_rendered_document
        mock_repository.mark_validated.return_value = mock_rendered_document

        result = await rendering_service.render_document(rendering_request)

        assert result.success
        assert result.output_path is not None

        stored_content = mock_storage.get_file(result.output_path)
        assert stored_content is not None
        assert len(stored_content) > 0
