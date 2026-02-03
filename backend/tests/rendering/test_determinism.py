import io

from docx import Document

from backend.app.domains.rendering.engine import DocumentRenderer
from backend.app.domains.rendering.validator import RenderedDocumentValidator


class TestSameInputProducesIdenticalOutput:
    def test_same_structure_produces_identical_content_hash(
        self,
        document_renderer: DocumentRenderer,
        simple_assembled_structure: dict,
    ):
        content1, _ = document_renderer.render(simple_assembled_structure)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(simple_assembled_structure)

        from backend.app.domains.rendering.schemas import compute_file_hash

        hash1 = compute_file_hash(content1)
        hash2 = compute_file_hash(content2)

        assert hash1 == hash2

    def test_same_structure_produces_same_block_count(
        self,
        document_renderer: DocumentRenderer,
        complex_assembled_structure: dict,
    ):
        content1, stats1 = document_renderer.render(complex_assembled_structure)

        renderer2 = DocumentRenderer()
        content2, stats2 = renderer2.render(complex_assembled_structure)

        assert stats1.paragraphs == stats2.paragraphs
        assert stats1.headings == stats2.headings
        assert stats1.tables == stats2.tables
        assert stats1.lists == stats2.lists
        assert stats1.total_blocks == stats2.total_blocks

    def test_same_structure_produces_same_text_content(
        self,
        document_renderer: DocumentRenderer,
        complex_assembled_structure: dict,
    ):
        content1, _ = document_renderer.render(complex_assembled_structure)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(complex_assembled_structure)

        doc1 = Document(io.BytesIO(content1))
        doc2 = Document(io.BytesIO(content2))

        texts1 = [p.text for p in doc1.paragraphs]
        texts2 = [p.text for p in doc2.paragraphs]

        assert texts1 == texts2

    def test_same_table_structure_is_identical(
        self,
        document_renderer: DocumentRenderer,
        sample_table_block: dict,
    ):
        structure = {
            "blocks": [sample_table_block],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content1, _ = document_renderer.render(structure)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(structure)

        doc1 = Document(io.BytesIO(content1))
        doc2 = Document(io.BytesIO(content2))

        assert len(doc1.tables) == len(doc2.tables)

        for t1, t2 in zip(doc1.tables, doc2.tables):
            assert len(t1.rows) == len(t2.rows)
            for r1, r2 in zip(t1.rows, t2.rows):
                for c1, c2 in zip(r1.cells, r2.cells):
                    assert c1.text == c2.text


class TestReRenderingDoesNotCreateDivergentFiles:
    def test_multiple_renders_produce_identical_validation(
        self,
        document_renderer: DocumentRenderer,
        document_validator: RenderedDocumentValidator,
        complex_assembled_structure: dict,
    ):
        content1, _ = document_renderer.render(complex_assembled_structure)
        result1 = document_validator.validate(content1)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(complex_assembled_structure)
        result2 = document_validator.validate(content2)

        assert result1.is_valid == result2.is_valid
        assert result1.block_count == result2.block_count
        assert result1.paragraph_count == result2.paragraph_count
        assert result1.table_count == result2.table_count

    def test_validator_determinism_check_passes(
        self,
        document_renderer: DocumentRenderer,
        document_validator: RenderedDocumentValidator,
        simple_assembled_structure: dict,
    ):
        content1, _ = document_renderer.render(simple_assembled_structure)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(simple_assembled_structure)

        is_deterministic, message = document_validator.validate_determinism(content1, content2)

        assert is_deterministic

    def test_determinism_with_complex_document(
        self,
        document_renderer: DocumentRenderer,
        document_validator: RenderedDocumentValidator,
        complex_assembled_structure: dict,
    ):
        content1, _ = document_renderer.render(complex_assembled_structure)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(complex_assembled_structure)

        is_deterministic, message = document_validator.validate_determinism(content1, content2)

        assert is_deterministic

    def test_ten_consecutive_renders_are_identical(
        self,
        simple_assembled_structure: dict,
    ):
        from backend.app.domains.rendering.schemas import compute_file_hash

        hashes = []
        for _ in range(10):
            renderer = DocumentRenderer()
            content, _ = renderer.render(simple_assembled_structure)
            hashes.append(compute_file_hash(content))

        assert len(set(hashes)) == 1


class TestDifferentInputProducesDifferentOutput:
    def test_different_content_produces_different_hash(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure1 = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [{"text": "Content A"}],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        structure2 = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [{"text": "Content B"}],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content1, _ = document_renderer.render(structure1)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(structure2)

        from backend.app.domains.rendering.schemas import compute_file_hash

        hash1 = compute_file_hash(content1)
        hash2 = compute_file_hash(content2)

        assert hash1 != hash2

    def test_different_block_count_produces_different_hash(
        self,
        document_renderer: DocumentRenderer,
    ):
        structure1 = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [{"text": "Only one paragraph"}],
                }
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        structure2 = {
            "blocks": [
                {
                    "block_type": "paragraph",
                    "block_id": "blk_001",
                    "sequence": 0,
                    "runs": [{"text": "First paragraph"}],
                },
                {
                    "block_type": "paragraph",
                    "block_id": "blk_002",
                    "sequence": 1,
                    "runs": [{"text": "Second paragraph"}],
                },
            ],
            "metadata": {},
            "headers": [],
            "footers": [],
        }

        content1, _ = document_renderer.render(structure1)

        renderer2 = DocumentRenderer()
        content2, _ = renderer2.render(structure2)

        from backend.app.domains.rendering.schemas import compute_file_hash

        hash1 = compute_file_hash(content1)
        hash2 = compute_file_hash(content2)

        assert hash1 != hash2
